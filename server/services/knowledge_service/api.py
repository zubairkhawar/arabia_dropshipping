from datetime import datetime
import base64
import hashlib
from io import BytesIO
import logging
import re
from typing import Any, List, Optional, Dict, Tuple

import httpx
from fastapi import APIRouter, Depends, status, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import get_db
from models import KnowledgeSource


router = APIRouter()
logger = logging.getLogger(__name__)
MAX_EXCEL_ROWS_PER_SHEET = 5000


class KnowledgeSourceIn(BaseModel):
    tenant_id: int
    name: str
    type: str  # file | url | api
    url: str | None = None
    metadata: dict | None = None


class KnowledgeSourceOut(BaseModel):
    id: int
    tenant_id: int
    name: str
    type: str
    url: Optional[str]
    status: str
    chunk_count: int
    metadata: Optional[Dict[str, Any]]
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


def _to_knowledge_source_out(row: KnowledgeSource) -> KnowledgeSourceOut:
    return KnowledgeSourceOut(
        id=row.id,
        tenant_id=row.tenant_id,
        name=row.name,
        type=row.type,
        url=row.url,
        status=row.status,
        chunk_count=row.chunk_count,
        metadata=row.knowledge_metadata or {},
        created_at=row.created_at,
        updated_at=row.updated_at,
    )


def _needs_refresh(row: KnowledgeSource) -> bool:
    if row.type not in {"url", "api"}:
        return False
    md = row.knowledge_metadata or {}
    hours_raw = md.get("refresh_interval_hours")
    try:
        hours = int(hours_raw) if hours_raw is not None else 24
    except Exception:
        hours = 24
    if hours <= 0:
        return False
    last = str(md.get("last_fetched_at") or "").strip()
    if not last:
        return True
    try:
        dt = datetime.fromisoformat(last)
    except Exception:
        return True
    return (datetime.utcnow() - dt).total_seconds() >= hours * 3600


def _clean_text(raw: str) -> str:
    # Best-effort HTML/text cleanup without external parser deps.
    txt = re.sub(r"(?is)<script.*?>.*?</script>", " ", raw)
    txt = re.sub(r"(?is)<style.*?>.*?</style>", " ", txt)
    txt = re.sub(r"(?is)<[^>]+>", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _normalize_plain_extraction(raw: str) -> str:
    """
    Normalize PDF/plain text without collapsing newlines into spaces.
    _clean_text destroys line breaks and makes PDFs one blob with no clause boundaries.
    """
    t = (raw or "").replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"[ \t\f\v]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def _hard_chunk_slices(s: str, size: int, overlap: int) -> List[str]:
    """Fixed-size windows with overlap (used when merge tail still exceeds size)."""
    if not s:
        return []
    ov = max(0, min(overlap, size // 2))
    out: List[str] = []
    start = 0
    L = len(s)
    while start < L:
        end = min(L, start + size)
        piece = s[start:end].strip()
        if piece:
            out.append(piece)
        if end >= L:
            break
        nxt = max(start + 1, end - ov)
        if nxt <= start:
            nxt = end
        start = nxt
    return out


def _split_sentences(text: str) -> List[str]:
    """
    Split into units for chunking. PDFs often lack '. ' boundaries; keep line breaks.
    """
    raw = (text or "").replace("\r\n", "\n").strip()
    if not raw:
        return []
    pieces: List[str] = []
    for block in re.split(r"\n{2,}", raw):
        block = block.strip()
        if not block:
            continue
        parts = re.split(r"(?<=[.!?。！？])\s+|\n+", block)
        for p in parts:
            p = p.strip()
            if p:
                pieces.append(p)
    return pieces if pieces else ([raw] if raw else [])


def _chunk_text(
    text: str,
    chunk_size: int = 900,
    overlap: int = 150,
    *,
    source_name: str = "unknown",
    page: Optional[int] = None,
    start_index: int = 0,
) -> List[Dict[str, Any]]:
    if not text:
        return []
    size = max(300, chunk_size)
    ov = max(0, min(overlap, size // 2))
    normalized = re.sub(r"\n{3,}", "\n\n", text).strip()
    paragraphs = [p.strip() for p in normalized.split("\n\n") if p.strip()]
    units: List[str] = []
    for para in paragraphs if paragraphs else [normalized]:
        sentences = _split_sentences(para)
        if sentences:
            units.extend(sentences)
        elif para:
            units.append(para)

    windows: List[str] = []
    current = ""
    for unit in units:
        candidate = (f"{current} {unit}".strip() if current else unit).strip()
        if len(candidate) <= size:
            current = candidate
            continue
        if current:
            windows.append(current)
            merged = (current[-ov:] + " " + unit).strip() if ov > 0 else unit
            if len(merged) > size:
                windows.extend(_hard_chunk_slices(merged, size, ov))
                current = ""
            else:
                current = merged
        else:
            windows.extend(_hard_chunk_slices(unit, size, ov))
            current = ""
    if current:
        if len(current) > size:
            windows.extend(_hard_chunk_slices(current, size, ov))
        else:
            windows.append(current)

    out: List[Dict[str, Any]] = []
    cursor = 0
    for idx, chunk_text in enumerate([w for w in windows if w], start=start_index):
        start_pos = max(0, normalized.find(chunk_text[:40], cursor))
        end_pos = start_pos + len(chunk_text)
        cursor = max(cursor, end_pos - ov)
        row: Dict[str, Any] = {
            "text": chunk_text,
            "index": idx,
            "source": source_name,
            "start_char": start_pos,
            "end_char": end_pos,
        }
        if page is not None:
            row["page"] = page
        out.append(row)
    return out


def _extract_pdf_page_text(page: Any, page_idx: int, src_name: str) -> str:
    """Best-effort page text; try default extract then layout mode when pypdf supports it."""
    chunks: List[str] = []
    try:
        chunks.append(page.extract_text() or "")
    except Exception as e:
        logger.warning("KB PDF %s p%d extract_text failed: %s", src_name, page_idx, e)
    try:
        chunks.append(page.extract_text(extraction_mode="layout") or "")
    except TypeError:
        pass
    except Exception as e:
        logger.debug("KB PDF %s p%d layout extract: %s", src_name, page_idx, e)
    for raw in chunks:
        s = (raw or "").strip()
        if s:
            return s
    return ""


def _ingest_pdf_blob(
    blob: bytes,
    src: KnowledgeSource,
    md: Dict[str, Any],
) -> Tuple[Optional[Tuple[str, int, Dict[str, Any]]], str]:
    """
    Build chunk objects from PDF bytes.
    Returns ((status, chunk_count, md), "") when ingestion finishes successfully.
    Returns (None, continuation_text) when the generic `text` path should run
    (e.g. joined text after failed per-page chunking), or (None, "") after error.
    """
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(BytesIO(blob))
        pages: List[str] = []
        for i, page in enumerate(reader.pages):
            pt = _extract_pdf_page_text(page, i + 1, src.name)
            if not pt:
                logger.warning("KB PDF page %s p%d: extracted no text", src.name, i + 1)
            pages.append(pt)
        raw_chars = sum(len(p) for p in pages)
        logger.info(
            "KB PDF extract %s: %d pages, %d raw chars from pypdf",
            src.name,
            len(pages),
            raw_chars,
        )
        page_chunks: List[Dict[str, Any]] = []
        chunk_idx = 0
        for page_idx, page_text in enumerate(pages, start=1):
            if not page_text:
                continue
            normalized_page = _normalize_plain_extraction(page_text)
            if not normalized_page:
                continue
            c = _chunk_text(
                normalized_page,
                source_name=src.name,
                page=page_idx,
                start_index=chunk_idx,
            )
            logger.debug(
                "KB PDF page %s p%d: %d in chars → %d chunks",
                src.name,
                page_idx,
                len(normalized_page),
                len(c),
            )
            page_chunks.extend(c)
            chunk_idx += len(c)
        joined_preview = "\n\n".join(_normalize_plain_extraction(p) for p in pages if p).strip()
        logger.info(
            "KB PDF %s after normalize: joined len=%d; page_chunks=%d",
            src.name,
            len(joined_preview),
            len(page_chunks),
        )
        if page_chunks:
            preview0 = (page_chunks[0].get("text") or "")[:500]
            logger.info(
                "PDF extracted %d raw chars, created %d chunks; first chunk len=%d",
                raw_chars,
                len(page_chunks),
                len(page_chunks[0].get("text") or ""),
            )
            logger.info("KB PDF first chunk preview (500): %r", preview0)
            md["chunks"] = page_chunks
            md["chunk_schema"] = "v2_object"
            md["parse_method"] = "pdf_pypdf_page_text"
            md["content_preview"] = preview0
            joined = "\n".join([str(x.get("text") or "") for x in page_chunks])
            md["content_hash"] = hashlib.sha256(joined.encode("utf-8", errors="ignore")).hexdigest()
            return ("ready", len(page_chunks), md), ""
        if joined_preview:
            joined_chunks = _chunk_text(joined_preview, source_name=src.name)
            if joined_chunks:
                logger.info(
                    "KB PDF ingest %s (joined fallback): %d chunks",
                    src.name,
                    len(joined_chunks),
                )
                logger.info(
                    "PDF extracted %d raw chars, created %d chunks (joined)",
                    raw_chars,
                    len(joined_chunks),
                )
                md["chunks"] = joined_chunks
                md["chunk_schema"] = "v2_object"
                md["parse_method"] = "pdf_pypdf_joined_text"
                md["content_preview"] = (joined_chunks[0].get("text") or "")[:500]
                jh = "\n".join(str(x.get("text") or "") for x in joined_chunks)
                md["content_hash"] = hashlib.sha256(jh.encode("utf-8", errors="ignore")).hexdigest()
                return ("ready", len(joined_chunks), md), ""
            logger.warning(
                "KB PDF joined chunking produced 0 chunks for %s (%d chars)",
                src.name,
                len(joined_preview),
            )
            return None, joined_preview
        return None, ""
    except Exception as e:
        logger.error("PDF extraction failed for %s: %s", src.name, e)
        md["last_error"] = f"pdf_extract_failed:{e!s}"[:220]
        return None, ""


async def _fetch_url_text(
    url: str,
    *,
    etag: Optional[str] = None,
    last_modified: Optional[str] = None,
) -> Tuple[str, Dict[str, str], bool]:
    req_headers: Dict[str, str] = {}
    if etag:
        req_headers["If-None-Match"] = etag
    if last_modified:
        req_headers["If-Modified-Since"] = last_modified
    async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as client:
        resp = await client.get(url, headers=req_headers or None)
        if resp.status_code == 304:
            return "", {}, True
        resp.raise_for_status()
        content_type = (resp.headers.get("content-type") or "").lower()
        raw = resp.text if "text" in content_type or "html" in content_type else resp.text
        fetch_headers = {
            "etag": resp.headers.get("etag") or "",
            "last_modified": resp.headers.get("last-modified") or "",
        }
    return _clean_text(raw), fetch_headers, False


def _build_api_source_text(src: KnowledgeSource) -> str:
    md = src.knowledge_metadata or {}
    schema_notes = str(md.get("schema_notes") or "").strip()
    auth_type = str(md.get("auth_type") or "none").strip()
    refresh = md.get("refresh_interval_hours")
    headers = md.get("headers") or {}
    if isinstance(headers, dict):
        header_keys = ", ".join(sorted(str(k) for k in headers.keys()))
    else:
        header_keys = ""
    lines = [
        f"API source: {src.name}",
        f"Base URL: {src.url or 'N/A'}",
        f"Auth type: {auth_type or 'none'}",
        f"Refresh interval hours: {refresh if refresh is not None else 'N/A'}",
    ]
    if header_keys:
        lines.append(f"Header keys: {header_keys}")
    if schema_notes:
        lines.append(f"Schema notes: {schema_notes}")
    return "\n".join(lines)


def _extract_docx_text(blob: bytes, source_name: str) -> str:
    try:
        import docx  # type: ignore

        doc = docx.Document(BytesIO(blob))
        parts: List[str] = []
        for para in doc.paragraphs:
            txt = (para.text or "").strip()
            if txt:
                parts.append(txt)
        for table in doc.tables:
            for row in table.rows:
                row_vals: List[str] = []
                for cell in row.cells:
                    val = (cell.text or "").strip()
                    if val:
                        row_vals.append(val)
                if row_vals:
                    parts.append(" | ".join(row_vals))
        return "\n\n".join(parts).strip()
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", source_name, exc)
        return ""


def _extract_excel_text(blob: bytes, source_name: str) -> str:
    try:
        import pandas as pd  # type: ignore

        book = pd.ExcelFile(BytesIO(blob))
        out: List[str] = []
        for sheet in book.sheet_names:
            try:
                df = pd.read_excel(book, sheet_name=sheet)
            except Exception as exc:
                logger.warning("Excel sheet parse failed for %s (%s): %s", source_name, sheet, exc)
                continue
            if df is None or df.empty:
                continue
            out.append(f"Sheet: {sheet}")
            # Keep bounded for very large files
            if len(df) > MAX_EXCEL_ROWS_PER_SHEET:
                df = df.head(MAX_EXCEL_ROWS_PER_SHEET)
            for _, row in df.iterrows():
                row_parts: List[str] = []
                for col, val in row.items():
                    if pd.notna(val):
                        row_parts.append(f"{col}: {val}")
                if row_parts:
                    out.append(" | ".join(row_parts))
            out.append("")
        return "\n".join(out).strip()
    except Exception as exc:
        logger.error("Excel extraction failed for %s: %s", source_name, exc)
        return ""


async def _ingest_source_content(src: KnowledgeSource) -> Tuple[str, int, Dict[str, Any]]:
    """
    Returns: status, chunk_count, metadata_patch
    """
    md = dict(src.knowledge_metadata or {})
    md.pop("last_error", None)
    md["last_fetched_at"] = datetime.utcnow().isoformat()

    if src.type == "url":
        if not src.url:
            md["last_error"] = "URL source missing url"
            return "error", 0, md
        text, fetch_headers, not_modified = await _fetch_url_text(
            src.url,
            etag=str(md.get("etag") or "").strip() or None,
            last_modified=str(md.get("last_modified") or "").strip() or None,
        )
        if not_modified:
            existing_chunks = md.get("chunks")
            if isinstance(existing_chunks, list) and existing_chunks:
                return "ready", len(existing_chunks), md
        if not text:
            md["last_error"] = "No readable content fetched from URL"
            return "error", 0, md
        chunks = _chunk_text(text, source_name=src.name)
        md["chunks"] = chunks
        md["chunk_schema"] = "v2_object"
        md["content_preview"] = text[:500]
        md["content_hash"] = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
        if fetch_headers.get("etag"):
            md["etag"] = fetch_headers["etag"]
        if fetch_headers.get("last_modified"):
            md["last_modified"] = fetch_headers["last_modified"]
        md["parse_method"] = "url_html_clean"
        return "ready", len(chunks), md

    if src.type == "api":
        text = _build_api_source_text(src)
        chunks = _chunk_text(text, chunk_size=700, overlap=80, source_name=src.name)
        md["chunks"] = chunks
        md["chunk_schema"] = "v2_object"
        md["content_preview"] = text[:500]
        md["content_hash"] = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
        md["parse_method"] = "api_metadata"
        return "ready", len(chunks), md

    # file source ingestion
    text = str(md.get("file_text") or "").strip()
    if not text:
        b64 = str(md.get("file_data_base64") or "").strip()
        mime = str(md.get("mime_type") or "").lower().strip()
        if b64:
            try:
                blob = base64.b64decode(b64)
                filename = str(md.get("filename") or src.name or "").lower()
                # Browsers often send application/octet-stream for PDFs; trust extension too.
                is_pdf = (
                    mime == "application/pdf"
                    or mime == "application/x-pdf"
                    or filename.endswith(".pdf")
                )
                if mime.startswith("text/") or mime in {
                    "application/json",
                    "application/csv",
                    "text/csv",
                }:
                    text = blob.decode("utf-8", errors="ignore")
                elif is_pdf:
                    pdf_done, text = _ingest_pdf_blob(blob, src, md)
                    if pdf_done is not None:
                        return pdf_done
                elif mime in {
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }:
                    text = _extract_docx_text(blob, src.name)
                    if not text:
                        md["last_error"] = "docx_extract_failed"
                elif mime in {
                    "application/vnd.ms-excel",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                }:
                    text = _extract_excel_text(blob, src.name)
                    if not text:
                        md["last_error"] = "excel_extract_failed"
                else:
                    # Fallback by extension when browsers provide unknown MIME.
                    filename = str(md.get("filename") or src.name or "").lower()
                    if filename.endswith(".docx"):
                        text = _extract_docx_text(blob, src.name)
                        if not text:
                            md["last_error"] = "docx_extract_failed"
                    elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                        text = _extract_excel_text(blob, src.name)
                        if not text:
                            md["last_error"] = "excel_extract_failed"
            except Exception:
                text = ""

    if text:
        mime_lower = str(md.get("mime_type") or "").lower().strip()
        fname = str(md.get("filename") or src.name or "").lower()
        if mime_lower == "application/pdf" or fname.endswith(".pdf"):
            cleaned = _normalize_plain_extraction(text)
            parse_method = "file_pdf_plain_normalize"
        else:
            cleaned = _clean_text(text)
            parse_method = "file_text_clean"
        chunks = _chunk_text(cleaned, source_name=src.name)
        md["chunks"] = chunks
        md["chunk_schema"] = "v2_object"
        md["content_preview"] = cleaned[:500]
        md["content_hash"] = hashlib.sha256(cleaned.encode("utf-8", errors="ignore")).hexdigest()
        md["parse_method"] = parse_method
        logger.info(
            "KB file ingest %s: %d chars → %d chunks (%s)",
            src.name,
            len(cleaned),
            len(chunks),
            parse_method,
        )
        return "ready", len(chunks), md

    filename = str(md.get("filename") or src.name or "file").strip()
    md["chunks"] = [
        {
            "text": f"File source connected: {filename}",
            "index": 0,
            "source": src.name,
            "start_char": 0,
            "end_char": len(filename),
        }
    ]
    md["chunk_schema"] = "v2_object"
    md["content_preview"] = f"File source connected: {filename}"
    md["parse_method"] = "file_placeholder"
    return "ready", 1, md


@router.get("/sources", response_model=List[KnowledgeSourceOut])
async def list_sources(tenant_id: int, db: Session = Depends(get_db)):
    """
    List knowledge sources for a tenant.
    """
    rows = (
        db.query(KnowledgeSource)
        .filter(KnowledgeSource.tenant_id == tenant_id)
        .order_by(KnowledgeSource.created_at.desc())
        .all()
    )
    return [_to_knowledge_source_out(row) for row in rows]


@router.post(
    "/sources",
    response_model=KnowledgeSourceOut,
    status_code=status.HTTP_201_CREATED,
)
async def create_source(payload: KnowledgeSourceIn, db: Session = Depends(get_db)):
    """
    Create a knowledge source.
    Indexing is stubbed; status starts as 'indexing'.
    """
    src = KnowledgeSource(
        tenant_id=payload.tenant_id,
        name=payload.name,
        type=payload.type,
        url=payload.url,
        status="indexing",
        chunk_count=0,
        knowledge_metadata=payload.metadata or {},
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(src)
    db.commit()
    db.refresh(src)

    # Practical synchronous ingestion for URL/API so KB becomes usable immediately.
    status_value, chunk_count, md_patch = await _ingest_source_content(src)
    src.status = status_value
    src.chunk_count = chunk_count
    src.knowledge_metadata = md_patch
    src.updated_at = datetime.utcnow()
    db.add(src)
    db.commit()
    db.refresh(src)

    return _to_knowledge_source_out(src)


@router.post("/sources/{source_id}/reindex", response_model=KnowledgeSourceOut)
async def reindex_source(source_id: int, db: Session = Depends(get_db)):
    src = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not src:
        raise HTTPException(status_code=404, detail="Knowledge source not found")
    src.status = "indexing"
    src.updated_at = datetime.utcnow()
    db.add(src)
    db.commit()
    db.refresh(src)

    status_value, chunk_count, md_patch = await _ingest_source_content(src)
    src.status = status_value
    src.chunk_count = chunk_count
    src.knowledge_metadata = md_patch
    src.updated_at = datetime.utcnow()
    db.add(src)
    db.commit()
    db.refresh(src)
    return _to_knowledge_source_out(src)


@router.post("/sources/reindex-stale", response_model=List[KnowledgeSourceOut])
async def reindex_stale_sources(tenant_id: int, db: Session = Depends(get_db)):
    """
    Reindex URL/API sources whose refresh interval has elapsed.
    """
    rows = (
        db.query(KnowledgeSource)
        .filter(KnowledgeSource.tenant_id == tenant_id)
        .order_by(KnowledgeSource.updated_at.asc())
        .all()
    )
    updated: List[KnowledgeSourceOut] = []
    for src in rows:
        if not _needs_refresh(src):
            continue
        src.status = "indexing"
        src.updated_at = datetime.utcnow()
        db.add(src)
        db.commit()
        db.refresh(src)
        status_value, chunk_count, md_patch = await _ingest_source_content(src)
        src.status = status_value
        src.chunk_count = chunk_count
        src.knowledge_metadata = md_patch
        src.updated_at = datetime.utcnow()
        db.add(src)
        db.commit()
        db.refresh(src)
        updated.append(_to_knowledge_source_out(src))
    return updated


@router.delete("/sources/{source_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_source(source_id: int, db: Session = Depends(get_db)):
    """
    Delete a knowledge source.
    """
    src = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
    if not src:
        return
    db.delete(src)
    db.commit()
    # TODO: remove from vector store when plugged in.
    return

